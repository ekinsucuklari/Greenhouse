from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_info_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Project Theme", "Cyber-Physical Greenhouse Simulation in Unity"),
        ("Current Phase", "Minimum Working Prototype + UI Refactor"),
        ("Primary Toolchain", "Unity 6, C#, TextMeshPro"),
        ("Submission Type", "First Progress Report"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v


def add_bdd_diagram(doc):
    add_heading(doc, "3.1 SysML BDD - Greenhouse System Structure", level=2)
    add_body(doc, "The block structure is shown below without overlapping connectors:")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    for col in table.columns:
        col.width = Inches(2.0)

    table.cell(0, 1).text = "GreenhouseSystem"
    table.cell(1, 0).text = "SimulationClock"
    table.cell(1, 1).text = "GreenhouseManager"
    table.cell(1, 2).text = "EnvironmentPhysics"
    table.cell(2, 0).text = "RuleBasedController"
    table.cell(2, 1).text = "SoilModel"
    table.cell(2, 2).text = "PlantGrowthModel"
    table.cell(3, 1).text = "Dashboard (Metrics + Actuators)"

    add_body(
        doc,
        "Key relations: GreenhouseManager composes AirState, SoilState, PlantState, and OutdoorState; "
        "it references SimulationClock, EnvironmentPhysics, SoilModel, and RuleBasedController.",
    )


def add_ibd_diagram(doc):
    add_heading(doc, "3.2 SysML IBD - Internal Data and Control Flow", level=2)
    add_body(doc, "Signal-level flow is listed in a clean matrix format:")

    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    headers = ("From", "To", "Data / Signal")
    for j, h in enumerate(headers):
        table.cell(0, j).text = h

    rows = [
        ("SimulationClock", "GreenhouseManager", "dt, simTime, hourOfDay"),
        ("EnvironmentPhysics", "GreenhouseManager", "airState, outdoorState"),
        ("SoilModel", "GreenhouseManager", "soilState"),
        ("PlantGrowthModel", "GreenhouseManager", "plantState"),
        ("RuleBasedController", "GreenhouseManager", "actuator decisions"),
        ("GreenhouseManager", "DashboardMetricsRuntime", "selected metric values"),
        ("GreenhouseManager", "DashboardActuatorsRuntime", "actuator status + UI commands"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.cell(i, j).text = val


def add_activity_diagram(doc):
    add_heading(doc, "3.3 SysML Activity - Simulation Tick", level=2)
    add_body(doc, "Execution order per fixed-step cycle:")
    steps = [
        "Start",
        "Compute dt from SimulationClock",
        "Update outdoor state",
        "Update indoor air state",
        "Update soil and plant state",
        "Evaluate rule-based controller",
        "Update energy/logging modules",
        "Refresh dashboard values",
        "End",
    ]
    for i, step in enumerate(steps, start=1):
        add_body(doc, f"{i}. {step}")


def add_state_diagram(doc):
    add_heading(doc, "3.4 SysML State Machine - Grow Light Controller", level=2)
    add_body(doc, "State transitions are presented in transition-table format:")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    headers = ("Current State", "Condition", "Next State", "Action")
    for j, h in enumerate(headers):
        table.cell(0, j).text = h

    table.cell(1, 0).text = "LightOff"
    table.cell(1, 1).text = "air.lightLux < threshold_on"
    table.cell(1, 2).text = "LightOn"
    table.cell(1, 3).text = "Set growLightActive = true"

    table.cell(2, 0).text = "LightOn"
    table.cell(2, 1).text = "air.lightLux > threshold_off"
    table.cell(2, 2).text = "LightOff"
    table.cell(2, 3).text = "Set growLightActive = false"

    add_body(doc, "Constraint: threshold_on must be lower than threshold_off to prevent ON/OFF chatter.")


def build_docx(output_path: str):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "Greenhouse Digital Twin Project")
    add_title(doc, "First Progress Report")
    p = doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    add_info_table(doc)
    doc.add_page_break()

    add_heading(doc, "1. Introduction", level=1)
    add_body(
        doc,
        "This report summarizes the first implementation stage of the Greenhouse Digital Twin in Unity. "
        "The main objective has been to stabilize the simulation loop, centralize system state, and build a scalable UI layer "
        "for monitoring and control in real time.",
    )

    add_heading(doc, "2. Work Completed", level=1)
    add_body(
        doc,
        "2.1 Core orchestration: GreenhouseManager now serves as the single integration point for air, soil, plant, outdoor, and actuator states.",
    )
    add_body(
        doc,
        "2.2 Physics and time: SimulationClock, EnvironmentPhysics, and SoilModel are integrated under fixed-step updates.",
    )
    add_body(
        doc,
        "2.3 Rule-based control: a hysteresis-based RuleBasedController is connected to fan, heater, irrigation, mister, and grow light logic.",
    )
    add_body(
        doc,
        "2.4 Dashboard refactor: UI was migrated to prefab-driven runtime generation for both metrics and actuator controls.",
    )

    add_heading(doc, "3. SysML Artifacts", level=1)
    add_bdd_diagram(doc)
    add_ibd_diagram(doc)
    add_activity_diagram(doc)
    add_state_diagram(doc)

    add_heading(doc, "4. Challenges", level=1)
    add_body(
        doc,
        "State wiring complexity during scene refactors occasionally caused null references or stale UI values.",
    )
    add_body(
        doc,
        "At high time-scale settings, control behavior around thresholds became sensitive and required additional calibration.",
    )
    add_body(
        doc,
        "Layout consistency across metrics and actuators required redesign to keep readability stable on different resolutions.",
    )

    add_heading(doc, "5. Next Steps", level=1)
    next_steps = [
        "Finalize responsive dashboard spacing and typography standards.",
        "Format simulation time display in HH:mm style for clarity.",
        "Tune grow-light and related threshold parameters to reduce chatter.",
        "Add validation scenarios and collect baseline performance logs.",
        "Convert draft SysML artifacts into final polished diagram set for submission.",
    ]
    for i, item in enumerate(next_steps, start=1):
        add_body(doc, f"{i}. {item}")

    add_heading(doc, "6. Conclusion", level=1)
    add_body(
        doc,
        "The project now has a stable architectural backbone with centralized state management, an operational simulation pipeline, "
        "and a configurable runtime dashboard. The next phase will focus on calibration quality, clearer reporting evidence, "
        "and final SysML documentation quality.",
    )

    doc.save(output_path)


if __name__ == "__main__":
    build_docx("First_Progress_Report_Editable.docx")
